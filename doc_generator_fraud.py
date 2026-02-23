"""
DOC GENERATOR - FRAUD
doc_generator_fraud.py

Word document version of the fraud chargeback response.
Same function signature as chargeback_generator_fraud.generate_pdf.
"""

import os
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _download_image(url, filename):
    try:
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            temp_path = os.path.join(os.environ.get('TEMP', '/tmp'), filename)
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            return temp_path
    except Exception as e:
        print(f"Error downloading image: {e}")
    return None


def _add_section_heading(doc, text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
    return h


def _add_image_or_placeholder(doc, image_path, caption=None, max_width=6.0):
    if image_path and os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(max_width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(8)
            return True
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")

    placeholder = doc.add_paragraph(f"[ {caption or 'Screenshot not available'} ]")
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in placeholder.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    return False


def _add_transaction_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        label_run = cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)
        cells[1].paragraphs[0].add_run(str(value) if value else '')
    return table


def generate_doc(data, kyc_images, output_path, session_evidence=None, tenant_name=None,
                 screenshots=None, public_records_data=None, location_map_data=None):
    """
    Generate fraud chargeback Word document.

    Args:
        data: LLM response data
        kyc_images: Dict with KYC image URLs
        output_path: Where to save .docx
        session_evidence: Session evidence from database
        tenant_name: Tenant/store name
        screenshots: Dict with screenshot paths
        public_records_data: Dict with public records from Redis
        location_map_data: Dict with location map screenshot and analysis
    """
    screenshots = screenshots or {}

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    has_kyc = any([kyc_images.get('id_card'), kyc_images.get('selfie'), kyc_images.get('card')])
    subtitle_text = ("Formal Evidence Submission with KYC Verification" if has_kyc
                     else "Formal Evidence Submission")

    title = doc.add_heading('CHARGEBACK DISPUTE RESPONSE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    sub = doc.add_paragraph(subtitle_text)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    doc.add_paragraph()

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    chargeback_reason = data.get('chargeback_reason') or 'Not Specified'
    if isinstance(chargeback_reason, str):
        chargeback_reason = chargeback_reason.replace('_', ' ').title()
    reference = data.get('reference') or data.get('transaction_id', '')

    _add_transaction_table(doc, [
        ["Reference:", reference],
        ["Amount:", f"{amount_str} {data.get('currency', 'USD')}"],
        ["Transaction Date:", data.get('transaction_date', '')],
        ["Dispute Reason:", chargeback_reason],
    ])
    doc.add_paragraph()

    if data.get('opening_statement'):
        para = doc.add_paragraph()
        bold_run = para.add_run('SUMMARY: ')
        bold_run.bold = True
        bold_run.font.size = Pt(10)
        rest = para.add_run(data['opening_statement'])
        rest.font.size = Pt(10)
        doc.add_paragraph()

    section_number = 1

    # Build interaction history text from session evidence
    interaction_history_text = None
    if session_evidence:
        parts = []
        for key in ('session_evidence', 'location_evidence', 'device_evidence'):
            sub_data = session_evidence.get(key, {})
            if sub_data and sub_data.get('summary'):
                parts.append(sub_data['summary'])
        if parts:
            interaction_history_text = '\n\n'.join(parts)

    evidence_sections = [
        ("ORDER DETAILS", "order_details", True),
        ("PAYMENT VERIFICATION", "payment_proof", True),
        ("IDENTITY VERIFICATION", "identity_proof", True),
        ("KYC VERIFICATION", "kyc_proof", False),
        ("PUBLIC RECORDS VERIFICATION", "public_records_proof", False),
        ("SHIPPING VERIFICATION", "shipping_proof", True),
        ("IP LOCATION VERIFICATION", "location_proof", False),
        ("CUSTOMER INTERACTION HISTORY", "interaction_proof", True),
    ]

    for header, key, always_show in evidence_sections:
        proof_data = data.get(key)

        if key == "order_details":
            _add_section_heading(doc, f'{section_number}. {header}')
            order_text = (f"Order {reference} placed on {data.get('transaction_date', '')} "
                          f"totaling {amount_str} {data.get('currency', 'USD')}.")
            if proof_data:
                order_text = (proof_data.get('text', order_text) if isinstance(proof_data, dict)
                              else proof_data)
            doc.add_paragraph(order_text)
            _add_image_or_placeholder(doc, screenshots.get('order_screenshot'), 'Shopify Order Details')
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "payment_proof":
            _add_section_heading(doc, f'{section_number}. {header}')
            doc.add_paragraph(
                "The following payment details were captured and verified during the transaction:")
            _add_image_or_placeholder(doc, screenshots.get('card_details_screenshot'),
                                      'Payment Gateway Card Details')
            if screenshots.get('avs_screenshot'):
                doc.add_paragraph()
                if proof_data:
                    payment_text = (proof_data.get('text') if isinstance(proof_data, dict)
                                    else proof_data)
                    if payment_text:
                        doc.add_paragraph(payment_text)
                _add_image_or_placeholder(doc, screenshots.get('avs_screenshot'),
                                          'AVS & Payment Verification Details')
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "identity_proof":
            _add_section_heading(doc, f'{section_number}. {header}')
            identity_text = ("The following payment information was collected and verified "
                             "during the transaction:")
            if proof_data:
                identity_text = (proof_data.get('text', identity_text) if isinstance(proof_data, dict)
                                 else proof_data)
            doc.add_paragraph(identity_text)
            _add_image_or_placeholder(doc, screenshots.get('identity_screenshot'),
                                      'Payment Identity Information')
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "kyc_proof" and has_kyc:
            _add_section_heading(doc, f'{section_number}. {header}')
            kyc_text = ("Customer completed comprehensive identity verification including "
                        "government-issued ID validation and live facial recognition.")
            if proof_data:
                t = proof_data.get('text') if isinstance(proof_data, dict) else proof_data
                if t:
                    kyc_text = t
            doc.add_paragraph(kyc_text)

            id_path = (_download_image(kyc_images['id_card'], 'kyc_id.png')
                       if kyc_images.get('id_card') else None)
            selfie_path = (_download_image(kyc_images['selfie'], 'kyc_selfie.png')
                           if kyc_images.get('selfie') else None)
            card_path = (_download_image(kyc_images['card'], 'kyc_card.png')
                         if kyc_images.get('card') else None)

            for img_path, caption in [(id_path, 'Government-Issued ID'),
                                      (selfie_path, 'Live Selfie Verification'),
                                      (card_path, 'Payment Card Verification')]:
                if img_path:
                    _add_image_or_placeholder(doc, img_path, caption, max_width=2.5)
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "kyc_proof" and not has_kyc and not proof_data:
            continue

        if key == "public_records_proof" and proof_data:
            _add_section_heading(doc, f'{section_number}. {header}')
            pr_text = proof_data.get('text') if isinstance(proof_data, dict) else proof_data
            if pr_text:
                doc.add_paragraph(pr_text)
            if public_records_data:
                phone = public_records_data.get('_phone_number', '')
                if phone:
                    doc.add_paragraph(
                        f"The following public records are associated with the phone number {phone}, "
                        "which was provided by the customer when placing this order:"
                    )
                fields = [(k, v) for k, v in public_records_data.items() if not k.startswith('_')]
                if phone:
                    fields.insert(0, ('Phone Number', phone))
                if fields:
                    pr_table = doc.add_table(rows=len(fields), cols=2)
                    pr_table.style = 'Table Grid'
                    for i, (k, v) in enumerate(fields):
                        pr_table.rows[i].cells[0].text = str(k).replace('_', ' ').title()
                        pr_table.rows[i].cells[1].text = str(v)
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "shipping_proof":
            _add_section_heading(doc, f'{section_number}. {header}')
            shipping_text = ("Please find below the Delivery Confirmation Receipt confirming "
                             "successful delivery of the order:")
            if proof_data:
                shipping_text = (proof_data.get('text', shipping_text) if isinstance(proof_data, dict)
                                 else proof_data)
            doc.add_paragraph(shipping_text)
            _add_image_or_placeholder(doc, screenshots.get('tracking_screenshot'),
                                      'Carrier Delivery Confirmation')
            if screenshots.get('tracking_url'):
                url_para = doc.add_paragraph(f"Tracking link: {screenshots['tracking_url']}")
                for run in url_para.runs:
                    run.font.size = Pt(9)
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "location_proof" and proof_data:
            _add_section_heading(doc, f'{section_number}. {header}')
            if location_map_data and location_map_data.get('analysis', {}).get('summary_text'):
                location_text = location_map_data['analysis']['summary_text']
            elif isinstance(proof_data, dict):
                location_text = proof_data.get('text')
            else:
                location_text = proof_data
            if location_text:
                doc.add_paragraph(location_text)
            _add_image_or_placeholder(doc, screenshots.get('location_screenshot'),
                                      'IP Location vs Billing/Shipping Address')
            doc.add_paragraph()
            section_number += 1
            continue

        if key == "interaction_proof":
            _add_section_heading(doc, f'{section_number}. {header}')
            if interaction_history_text:
                doc.add_paragraph(interaction_history_text)
            elif proof_data:
                text = (proof_data.get('text', 'Customer interaction and communication history.')
                        if isinstance(proof_data, dict) else proof_data)
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("Customer interaction and communication history.")
            ph = doc.add_paragraph('[ INSERT: Customer Interaction Screenshot ]')
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            section_number += 1
            continue

        if always_show and not proof_data:
            _add_section_heading(doc, f'{section_number}. {header}')
            doc.add_paragraph("Evidence documentation.")
            ph = doc.add_paragraph(f'[ INSERT: {header} Screenshot ]')
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            section_number += 1
            continue

        if proof_data:
            text = proof_data.get('text') if isinstance(proof_data, dict) else proof_data
            if text:
                _add_section_heading(doc, f'{section_number}. {header}')
                doc.add_paragraph(text)
                placeholder_text = (proof_data.get('proof_placeholder', f'{header} Screenshot')
                                    if isinstance(proof_data, dict) else f'{header} Screenshot')
                ph = doc.add_paragraph(f'[ INSERT: {placeholder_text} ]')
                ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                section_number += 1

    # Closing statement
    doc.add_paragraph()
    if data.get('closing_statement'):
        _add_section_heading(doc, 'CONCLUSION')
        para = doc.add_paragraph(data['closing_statement'])
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    # Footer branding
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by FUGU Chargeback Response System')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.save(output_path)
    print(f"DOC created: {output_path}")
    return output_path
