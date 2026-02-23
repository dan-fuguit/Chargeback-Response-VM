"""
DOC GENERATOR - CNP
doc_generator_cnp.py

Word document version of the Credit Not Processed chargeback response.
Same function signature as chargeback_generator_cnp.generate_pdf.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from return_policies import get_return_policy

RETURN_POLICIES_FOLDER = "return_policies"


def _get_return_policy_image(tenant_name):
    if not tenant_name:
        return None
    tenant_key = tenant_name.lower().strip()
    for ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
        image_path = os.path.join(RETURN_POLICIES_FOLDER, f"{tenant_key}{ext}")
        if os.path.exists(image_path):
            return image_path
    return None


def _add_section_heading(doc, text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
    return h


def _add_image_or_placeholder(doc, image_path, caption=None, max_width=6.0):
    if image_path and os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(max_width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(8)
            return True
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")

    placeholder = doc.add_paragraph(f"[ {caption or 'Screenshot not available'} ]")
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in placeholder.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    return False


def generate_doc(data, output_path, tenant_name=None, screenshots=None):
    """
    Generate CNP chargeback Word document.

    Args:
        data: LLM response data
        output_path: Where to save .docx
        tenant_name: Tenant/store name
        screenshots: Dict with screenshot paths
    """
    screenshots = screenshots or {}
    return_policy = get_return_policy(tenant_name)
    return_policy_image = _get_return_policy_image(tenant_name)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    title = doc.add_heading('CHARGEBACK DISPUTE RESPONSE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    sub = doc.add_paragraph("Credit / Refund Not Processed")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    doc.add_paragraph()

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    chargeback_reason = data.get('chargeback_reason') or 'Credit Not Processed'
    if isinstance(chargeback_reason, str):
        chargeback_reason = chargeback_reason.replace('_', ' ').title()
    reference = data.get('reference') or data.get('transaction_id', '')

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ("Reference:", reference),
        ("Amount:", f"{amount_str} {data.get('currency', 'USD')}"),
        ("Transaction Date:", data.get('transaction_date', '')),
        ("Dispute Reason:", chargeback_reason),
    ]):
        cells = table.rows[i].cells
        r = cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        cells[1].paragraphs[0].add_run(str(value) if value else '')
    doc.add_paragraph()

    # Opening statement
    if data.get('opening_statement'):
        for line in data['opening_statement'].split('\n'):
            if line.strip():
                if line.strip().startswith('∙'):
                    doc.add_paragraph(f"• {line.strip()[1:].strip()}")
                else:
                    doc.add_paragraph(line.strip())
        doc.add_paragraph()
    else:
        gender = data.get('customer_gender', 'they').lower()
        pronoun = 'she' if gender == 'female' else ('he' if gender == 'male' else 'they')
        doc.add_paragraph("Dear Issuer,")
        doc.add_paragraph(
            f"Please be informed that we wish to decline the chargeback with reason code {chargeback_reason}. "
            "No refund or credit was agreed upon or issued for this transaction. The order was fulfilled and "
            f"delivered to the client as confirmed by the shipping proof below. The cardholder did not follow "
            "the merchant's Returns & Exchanges Policy prior to initiating the chargeback."
        )
        doc.add_paragraph()
        for item in ["Order details", "Payment verification", "Shipping proof",
                     "Merchant's Returns & Exchanges Policy"]:
            doc.add_paragraph(f"• {item}")
        doc.add_paragraph()

    # Section 1: Order Details
    _add_section_heading(doc, "1. Order details")
    order_text = "Please find below the details of the order which was created on the merchant's web-site:"
    if data.get('order_details') and isinstance(data['order_details'], dict):
        order_text = data['order_details'].get('text', order_text)
    doc.add_paragraph(order_text)
    _add_image_or_placeholder(doc, screenshots.get('order_screenshot'), 'Shopify Order Details')
    doc.add_paragraph()

    # Section 2: Payment Verification
    _add_section_heading(doc, "2. Payment Verification")
    payment_proof = data.get('payment_proof')
    payment_text = None
    if payment_proof:
        payment_text = payment_proof.get('text') if isinstance(payment_proof, dict) else payment_proof
    if payment_text:
        doc.add_paragraph(payment_text)
    else:
        doc.add_paragraph(
            "The following payment details confirm that the transaction was successfully authorised and "
            "captured. No credit or refund was issued against this payment."
        )
    _add_image_or_placeholder(doc, screenshots.get('card_details_screenshot'),
                              'Payment Gateway Card Details')
    doc.add_paragraph()

    # Section 3: Shipping Proof
    _add_section_heading(doc, "3. Shipping proof")
    shipping_text = ("Please find below the Delivery Confirmation Receipt confirming "
                     "successful delivery of the order:")
    if data.get('shipping_proof') and isinstance(data['shipping_proof'], dict):
        shipping_text = data['shipping_proof'].get('text', shipping_text)
    doc.add_paragraph(shipping_text)
    _add_image_or_placeholder(doc, screenshots.get('tracking_screenshot'),
                              'Carrier Delivery Confirmation')
    doc.add_paragraph()

    # Section 4: Return/Exchange Policy
    _add_section_heading(doc, "4. Merchant's Return/Exchange Policy")
    doc.add_paragraph(return_policy["text"])
    if return_policy.get("url"):
        doc.add_paragraph("Merchant's Returns & Exchanges Policy can be found by the below link:")
        url_para = doc.add_paragraph(return_policy["url"])
        for run in url_para.runs:
            run.font.color.rgb = RGBColor(0x2b, 0x6c, 0xb0)
    doc.add_paragraph("Here is an extract from merchant's refund policy:")
    _add_image_or_placeholder(doc, return_policy_image, 'Return Policy Screenshot', max_width=6.0)

    # Summary
    doc.add_paragraph()
    _add_section_heading(doc, "Summary")
    if data.get('closing_statement'):
        para = doc.add_paragraph(data['closing_statement'])
        for run in para.runs:
            run.bold = True
    else:
        closing = doc.add_paragraph(
            "The transaction was successfully authorised, captured, and the order was delivered to the "
            "cardholder as confirmed by the shipping documentation above. No credit or refund was agreed "
            "upon or processed. The cardholder did not attempt to return the product in accordance with the "
            "merchant's Returns & Exchanges Policy. These facts confirm that this chargeback is unwarranted "
            "and should be cancelled."
        )
        for run in closing.runs:
            run.bold = True

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by FUGU Chargeback Response System')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.save(output_path)
    print(f"DOC created: {output_path}")
    return output_path
