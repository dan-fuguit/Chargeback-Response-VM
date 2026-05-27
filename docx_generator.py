"""
Word Document Generator for Chargeback Responses
Generates .docx files matching the same structure as the PDF generators.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from public_records import create_public_records_table_data


LOGO_PATH = "fugu_logo.png"

DARK_BLUE = RGBColor(0x1a, 0x36, 0x5d)
SECTION_BLUE = RGBColor(0x2c, 0x52, 0x82)
GRAY_TEXT = RGBColor(0x4a, 0x55, 0x68)
BODY_TEXT = RGBColor(0x2d, 0x37, 0x48)
LINK_BLUE = RGBColor(0x25, 0x63, 0xeb)
LIGHT_BG = RGBColor(0xf7, 0xfa, 0xfc)
BORDER_COLOR = RGBColor(0xe2, 0xe8, 0xf0)
GREEN_BG = RGBColor(0xf0, 0xff, 0xf4)
FUGU_BLUE = RGBColor(0x28, 0x73, 0xc2)


def _set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2c5282')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_thin_rule(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'e2e8f0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    p.paragraph_format.space_after = Pt(12)
    return p


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = SECTION_BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_conclusion_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_link_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = LINK_BLUE
    run.underline = True
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_transaction_table(doc, data):
    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"

    chargeback_reason = data.get('chargeback_reason') or 'Not Specified'
    if isinstance(chargeback_reason, str):
        chargeback_reason = chargeback_reason.replace('_', ' ').title()

    reference = data.get('reference') or data.get('transaction_id', '')

    rows = [
        ("Reference:", str(reference)),
        ("Amount:", f"{amount_str} {data.get('currency', 'USD')}"),
        ("Transaction Date:", str(data.get('transaction_date', ''))),
        ("Dispute Reason:", str(chargeback_reason)),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        cell_label.text = ''
        cell_value.text = ''

        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(9)
        run_l.font.color.rgb = GRAY_TEXT
        cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(9)
        run_v.font.color.rgb = BODY_TEXT

        _set_cell_shading(cell_label, 'f7fafc')
        _set_cell_shading(cell_value, 'f7fafc')

    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(5.1)

    doc.add_paragraph()
    return table


def _add_image(doc, image_path, caption=None, max_width=6.0):
    if image_path and os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(max_width))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if caption:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                run.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = GRAY_TEXT
            return True
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {caption or 'Image Placeholder'} ]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    return False


def _add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        if os.path.exists(LOGO_PATH):
            try:
                run = p.add_run()
                run.add_picture(LOGO_PATH, height=Inches(0.25))
            except:
                pass
            run2 = p.add_run("  Powered by FUGU")
        else:
            run2 = p.add_run("Powered by FUGU")
        run2.font.size = Pt(8)
        run2.font.color.rgb = FUGU_BLUE
        run2.bold = True


def _download_image(url, filename):
    try:
        import requests
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            temp_path = os.path.join(os.environ.get('TEMP', '/tmp'), filename)
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            return temp_path
    except Exception as e:
        print(f"Error downloading: {e}")
    return None


def _build_interaction_history_text(session_evidence):
    if not session_evidence:
        return None

    parts = []
    session_data = session_evidence.get('session_evidence', {})
    if session_data and session_data.get('summary'):
        parts.append(session_data['summary'])

    location_data = session_evidence.get('location_evidence', {})
    if location_data and location_data.get('summary'):
        parts.append(location_data['summary'])

    device_data = session_evidence.get('device_evidence', {})
    if device_data and device_data.get('summary'):
        parts.append(device_data['summary'])

    if parts:
        return '\n\n'.join(parts)
    return None


def _has_kyc_images(kyc_images):
    return any([kyc_images.get('id_card'), kyc_images.get('selfie'), kyc_images.get('card')])


# ============================================================================
# FRAUD generator
# ============================================================================

def generate_docx_fraud(data, kyc_images, output_path, session_evidence=None,
                        tenant_name=None, screenshots=None, public_records_data=None,
                        location_map_data=None):
    screenshots = screenshots or {}
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    has_kyc = _has_kyc_images(kyc_images)
    subtitle_text = "Formal Evidence Submission with KYC Verification" if has_kyc else "Formal Evidence Submission"

    _add_title(doc, "CHARGEBACK DISPUTE RESPONSE")
    _add_subtitle(doc, subtitle_text)
    _add_horizontal_rule(doc)

    _add_transaction_table(doc, data)

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    reference = data.get('reference') or data.get('transaction_id', '')

    if data.get('opening_statement'):
        _add_body_text(doc, "SUMMARY: " + data['opening_statement'])

    section_number = 1
    interaction_history_text = _build_interaction_history_text(session_evidence)

    evidence_sections = [
        ("ORDER DETAILS", "order_details", True),
        ("PAYMENT VERIFICATION", "payment_proof", True),
        ("IDENTITY VERIFICATION", "identity_proof", True),
        ("KYC VERIFICATION", "kyc_proof", False),
        ("PUBLIC RECORDS VERIFICATION", "public_records_proof", False),
        ("SHIPPING VERIFICATION", "shipping_proof", True),
        ("IP LOCATION VERIFICATION", "location_proof", False),
        ("CUSTOMER INTERACTION HISTORY", "interaction_proof", True),
    ]

    for header, key, always_show in evidence_sections:
        proof_data = data.get(key)

        if key == "order_details":
            _add_section_header(doc, f"{section_number}. {header}")
            order_text = f"Order {reference} placed on {data.get('transaction_date', '')} totaling {amount_str} {data.get('currency', 'USD')}."
            if proof_data:
                if isinstance(proof_data, dict):
                    order_text = proof_data.get('text', order_text)
                else:
                    order_text = proof_data
            _add_body_text(doc, order_text)
            _add_image(doc, screenshots.get('order_screenshot'), "Shopify Order Details")
            section_number += 1
            continue

        if key == "payment_proof":
            _add_section_header(doc, f"{section_number}. {header}")
            _add_body_text(doc, "The following payment details were captured and verified during the transaction:")
            _add_image(doc, screenshots.get('card_details_screenshot'), "Payment Gateway Card Details")

            if screenshots.get('avs_screenshot'):
                if proof_data:
                    payment_text = None
                    if isinstance(proof_data, dict):
                        payment_text = proof_data.get('text')
                    else:
                        payment_text = proof_data
                    if payment_text:
                        _add_body_text(doc, payment_text)
                _add_image(doc, screenshots.get('avs_screenshot'), "AVS & Payment Verification Details")
            section_number += 1
            continue

        if key == "identity_proof":
            _add_section_header(doc, f"{section_number}. {header}")
            identity_text = "The following payment information was collected and verified during the transaction:"
            if proof_data:
                if isinstance(proof_data, dict):
                    identity_text = proof_data.get('text', identity_text)
                else:
                    identity_text = proof_data
            _add_body_text(doc, identity_text)
            _add_image(doc, screenshots.get('identity_screenshot'), "Payment Identity Information")
            section_number += 1
            continue

        if key == "kyc_proof" and has_kyc:
            _add_section_header(doc, f"{section_number}. {header}")
            kyc_text = None
            if proof_data:
                if isinstance(proof_data, dict):
                    kyc_text = proof_data.get('text')
                else:
                    kyc_text = proof_data
            if not kyc_text:
                kyc_text = "Customer completed comprehensive identity verification including government-issued ID validation and live facial recognition."
            _add_body_text(doc, kyc_text)

            for img_key, label in [('id_card', 'Government-Issued ID'), ('selfie', 'Live Selfie Verification'), ('card', 'Payment Card Verification')]:
                if kyc_images.get(img_key):
                    img_path = _download_image(kyc_images[img_key], f'kyc_{img_key}.png')
                    if img_path:
                        _add_image(doc, img_path, label, max_width=3.0)
            section_number += 1
            continue

        if key == "kyc_proof" and not has_kyc and not proof_data:
            continue

        if key == "public_records_proof" and proof_data:
            _add_section_header(doc, f"{section_number}. {header}")
            pr_text = None
            if isinstance(proof_data, dict):
                pr_text = proof_data.get('text')
            else:
                pr_text = proof_data
            if pr_text:
                _add_body_text(doc, pr_text)

            if public_records_data:
                phone_number = public_records_data.get('_phone_number', '')
                if phone_number:
                    _add_body_text(doc, f"The following public records are associated with the phone number {phone_number}, which was provided by the customer when placing this order:")

                table_data = create_public_records_table_data(public_records_data)
                if phone_number:
                    table_data.insert(0, ["Phone Number:", phone_number])

                if table_data:
                    t = doc.add_table(rows=len(table_data), cols=2)
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, (label, value) in enumerate(table_data):
                        cell_l = t.cell(i, 0)
                        cell_v = t.cell(i, 1)
                        cell_l.text = ''
                        cell_v.text = ''
                        rl = cell_l.paragraphs[0].add_run(str(label))
                        rl.bold = True
                        rl.font.size = Pt(9)
                        rl.font.color.rgb = GRAY_TEXT
                        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        rv = cell_v.paragraphs[0].add_run(str(value))
                        rv.font.size = Pt(9)
                        rv.font.color.rgb = BODY_TEXT
                        _set_cell_shading(cell_l, 'f0fff4')
                        _set_cell_shading(cell_v, 'f0fff4')
                    t.columns[0].width = Inches(1.8)
                    t.columns[1].width = Inches(4.7)
            section_number += 1
            continue

        if key == "shipping_proof":
            _add_section_header(doc, f"{section_number}. {header}")
            shipping_text = "Please find below the Delivery Confirmation Receipt confirming successful delivery of the order:"
            if proof_data:
                if isinstance(proof_data, dict):
                    shipping_text = proof_data.get('text', shipping_text)
                else:
                    shipping_text = proof_data
            _add_body_text(doc, shipping_text)
            _add_image(doc, screenshots.get('tracking_screenshot'), "Carrier Delivery Confirmation")

            if screenshots.get('tracking_url'):
                _add_link_text(doc, f"Tracking link: {screenshots['tracking_url']}")
            section_number += 1
            continue

        if key == "location_proof" and proof_data:
            _add_section_header(doc, f"{section_number}. {header}")
            location_text = None
            if location_map_data and location_map_data.get('analysis', {}).get('summary_text'):
                location_text = location_map_data['analysis']['summary_text']
            elif isinstance(proof_data, dict):
                location_text = proof_data.get('text')
            else:
                location_text = proof_data
            if location_text:
                _add_body_text(doc, location_text)
            _add_image(doc, screenshots.get('location_screenshot'), "IP Location vs Billing/Shipping Address")
            section_number += 1
            continue

        if key == "interaction_proof":
            _add_section_header(doc, f"{section_number}. {header}")
            if interaction_history_text:
                _add_body_text(doc, interaction_history_text)
            elif proof_data:
                if isinstance(proof_data, dict):
                    text = proof_data.get('text', 'Customer interaction and communication history.')
                else:
                    text = proof_data
                _add_body_text(doc, text)
            else:
                _add_body_text(doc, "Customer interaction and communication history.")
            section_number += 1
            continue

        if always_show and not proof_data:
            _add_section_header(doc, f"{section_number}. {header}")
            _add_body_text(doc, "Evidence documentation.")
            section_number += 1
            continue

        if proof_data:
            if isinstance(proof_data, dict):
                text = proof_data.get('text')
            else:
                text = proof_data
            if text:
                _add_section_header(doc, f"{section_number}. {header}")
                _add_body_text(doc, text)
                section_number += 1

    _add_thin_rule(doc)
    if data.get('closing_statement'):
        _add_section_header(doc, "CONCLUSION")
        _add_conclusion_text(doc, data['closing_statement'])

    _add_footer(doc)
    doc.save(output_path)
    print(f"DOCX created: {output_path}")
    return output_path


# ============================================================================
# PNR generator
# ============================================================================

def generate_docx_pnr(data, output_path, tenant_name=None, screenshots=None):
    screenshots = screenshots or {}
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    _add_title(doc, "CHARGEBACK DISPUTE RESPONSE")
    _add_subtitle(doc, "Merchandise/Services Not Received - Delivery Confirmed")
    _add_horizontal_rule(doc)
    _add_transaction_table(doc, data)

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    reference = data.get('reference') or data.get('transaction_id', '')
    chargeback_reason = data.get('chargeback_reason') or 'Merchandise Not Received'

    if data.get('opening_statement'):
        for line in data['opening_statement'].split('\n'):
            if line.strip():
                _add_body_text(doc, line.strip())
    else:
        carrier = data.get('carrier', 'the carrier')
        _add_body_text(doc, "Dear Issuer,")
        _add_body_text(doc,
            f"We respectfully decline chargeback {reference} with reason code {chargeback_reason}. "
            f"The order was successfully delivered to the cardholder and delivery was confirmed by {carrier}. "
            f"This chargeback is therefore invalid."
        )

    # Section 1: Order Details
    _add_section_header(doc, "1. Order")
    order_text = f"Order {reference} placed on {data.get('transaction_date', '')} totaling {amount_str} {data.get('currency', 'USD')}."
    if data.get('order_details'):
        if isinstance(data['order_details'], dict):
            order_text = data['order_details'].get('text', order_text)
    _add_body_text(doc, order_text)
    _add_image(doc, screenshots.get('order_screenshot'), "Shopify Order Details")

    # Section 2: Card Details
    _add_section_header(doc, "2. Card Details")
    _add_body_text(doc, "The following card details were captured from the payment gateway:")
    _add_image(doc, screenshots.get('card_details_screenshot'), "Payment Gateway Card Details")

    # Section 3: Shipping Proof
    _add_section_header(doc, "3. Shipping proof")
    shipping_text = "Please find below the Delivery Confirmation Receipt confirming successful delivery of the order:"
    if data.get('shipping_proof'):
        if isinstance(data['shipping_proof'], dict):
            shipping_text = data['shipping_proof'].get('text', shipping_text)
    _add_body_text(doc, shipping_text)
    _add_image(doc, screenshots.get('tracking_screenshot'), "Carrier Delivery Confirmation")

    if screenshots.get('tracking_url'):
        _add_link_text(doc, f"Tracking link: {screenshots['tracking_url']}")

    # Closing
    _add_thin_rule(doc)
    _add_section_header(doc, "Summary")
    if data.get('closing_statement'):
        _add_conclusion_text(doc, data['closing_statement'])
    else:
        _add_conclusion_text(doc,
            "As shown in the delivery confirmation provided above, the package was successfully delivered "
            "to the cardholder's address. The merchant is not responsible for packages after confirmed delivery. "
            "These facts confirm that the product was received and the chargeback should be cancelled."
        )

    _add_footer(doc)
    doc.save(output_path)
    print(f"DOCX created: {output_path}")
    return output_path


# ============================================================================
# PNA generator
# ============================================================================

def generate_docx_pna(data, output_path, tenant_name=None, screenshots=None):
    screenshots = screenshots or {}
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    from return_policies import get_return_policy
    return_policy = get_return_policy(tenant_name)

    return_policy_image = None
    if tenant_name:
        tenant_key = tenant_name.lower().strip()
        for ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            path = os.path.join("return_policies", f"{tenant_key}{ext}")
            if os.path.exists(path):
                return_policy_image = path
                break

    _add_title(doc, "CHARGEBACK DISPUTE RESPONSE")
    _add_subtitle(doc, "Product Unacceptable / Quality Dispute")
    _add_horizontal_rule(doc)
    _add_transaction_table(doc, data)

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    reference = data.get('reference') or data.get('transaction_id', '')
    chargeback_reason = data.get('chargeback_reason') or 'Product Unacceptable'
    if isinstance(chargeback_reason, str):
        chargeback_reason = chargeback_reason.replace('_', ' ').title()

    if data.get('opening_statement'):
        for line in data['opening_statement'].split('\n'):
            if line.strip():
                if line.strip().startswith('∙'):
                    _add_bullet_text(doc, line.strip())
                else:
                    _add_body_text(doc, line.strip())
    else:
        customer_name = data.get('customer_name', 'the customer')
        gender = data.get('customer_gender', 'they').lower()
        pronoun = 'she' if gender == 'female' else ('he' if gender == 'male' else 'they')
        _add_body_text(doc, "Dear Issuer,")
        _add_body_text(doc,
            f"Please be informed that we wish to decline the chargeback with reason code {chargeback_reason}. "
            f"The order was successfully delivered to the client in good condition and {pronoun} never claimed "
            f"about the quality of the item or asked for a return, so the chargeback is invalid."
        )
        _add_bullet_text(doc, "∙ Order details")
        _add_bullet_text(doc, "∙ Payment proof")
        _add_bullet_text(doc, "∙ Shipping proof")
        _add_bullet_text(doc, "∙ Merchant's Returns & Exchanges Policy")

    # Section 1: Order Details
    _add_section_header(doc, "1. Order details")
    order_text = "Please find below the details of order which was created on the merchant's web-site:"
    if data.get('order_details'):
        if isinstance(data['order_details'], dict):
            order_text = data['order_details'].get('text', order_text)
    _add_body_text(doc, order_text)
    _add_image(doc, screenshots.get('order_screenshot'), "Shopify Order Details")

    # Section 2: Card Details
    _add_section_header(doc, "2. Card Details")
    _add_body_text(doc, "The following card details were captured from the payment gateway:")
    _add_image(doc, screenshots.get('card_details_screenshot'), "Payment Gateway Card Details")

    # Section 3: Shipping Proof
    _add_section_header(doc, "3. Shipping proof")
    shipping_text = "Please find below the Delivery Confirmation Receipt confirming successful delivery of the order:"
    if data.get('shipping_proof'):
        if isinstance(data['shipping_proof'], dict):
            shipping_text = data['shipping_proof'].get('text', shipping_text)
    _add_body_text(doc, shipping_text)
    _add_image(doc, screenshots.get('tracking_screenshot'), "Carrier Delivery Confirmation")

    # Section 4: Return Policy
    _add_section_header(doc, "4. Merchant's Return/Exchange Policy")
    _add_body_text(doc, return_policy["text"])
    if return_policy.get("url"):
        _add_body_text(doc, "Merchant's Returns & Exchanges Policy can be found by the below link:")
        _add_link_text(doc, return_policy['url'])
    _add_body_text(doc, "Here is an extract from merchant's refund policy:")
    _add_image(doc, return_policy_image, "Return Policy Screenshot")

    # Closing
    _add_thin_rule(doc)
    _add_section_header(doc, "Summary")
    if data.get('closing_statement'):
        _add_conclusion_text(doc, data['closing_statement'])
    else:
        _add_conclusion_text(doc,
            "All products are inspected and packaged for shipment prior to leaving the warehouse. "
            "The cardholder has never claimed about the quality of the item. The cardholder has never "
            "tried to return the product according to merchant's Returns & Exchanges Policy. These facts "
            "confirm that this is a clear case of buyer's remorse and so the chargeback should be cancelled."
        )

    _add_footer(doc)
    doc.save(output_path)
    print(f"DOCX created: {output_path}")
    return output_path


# ============================================================================
# CNP generator
# ============================================================================

def generate_docx_cnp(data, output_path, tenant_name=None, screenshots=None):
    screenshots = screenshots or {}
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    from return_policies import get_return_policy
    return_policy = get_return_policy(tenant_name)

    return_policy_image = None
    if tenant_name:
        tenant_key = tenant_name.lower().strip()
        for ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            path = os.path.join("return_policies", f"{tenant_key}{ext}")
            if os.path.exists(path):
                return_policy_image = path
                break

    _add_title(doc, "CHARGEBACK DISPUTE RESPONSE")
    _add_subtitle(doc, "Credit / Refund Not Processed")
    _add_horizontal_rule(doc)
    _add_transaction_table(doc, data)

    amount = data.get('amount', 0)
    amount_str = f"${amount:.2f}" if isinstance(amount, (int, float)) else f"${amount}"
    reference = data.get('reference') or data.get('transaction_id', '')
    chargeback_reason = data.get('chargeback_reason') or 'Credit Not Processed'
    if isinstance(chargeback_reason, str):
        chargeback_reason = chargeback_reason.replace('_', ' ').title()

    if data.get('opening_statement'):
        for line in data['opening_statement'].split('\n'):
            if line.strip():
                if line.strip().startswith('∙'):
                    _add_bullet_text(doc, line.strip())
                else:
                    _add_body_text(doc, line.strip())
    else:
        _add_body_text(doc, "Dear Issuer,")
        _add_body_text(doc,
            f"Please be informed that we wish to decline the chargeback with reason code {chargeback_reason}. "
            f"No refund or credit was agreed upon or issued for this transaction. The order was fulfilled and "
            f"delivered to the client as confirmed by the shipping proof below. The cardholder did not follow "
            f"the merchant's Returns & Exchanges Policy prior to initiating the chargeback."
        )
        _add_bullet_text(doc, "∙ Order details")
        _add_bullet_text(doc, "∙ Payment verification")
        _add_bullet_text(doc, "∙ Shipping proof")
        _add_bullet_text(doc, "∙ Merchant's Returns & Exchanges Policy")

    # Section 1: Order Details
    _add_section_header(doc, "1. Order details")
    order_text = "Please find below the details of the order which was created on the merchant's web-site:"
    if data.get('order_details'):
        if isinstance(data['order_details'], dict):
            order_text = data['order_details'].get('text', order_text)
    _add_body_text(doc, order_text)
    _add_image(doc, screenshots.get('order_screenshot'), "Shopify Order Details")

    # Section 2: Payment Verification
    _add_section_header(doc, "2. Payment Verification")
    payment_proof = data.get('payment_proof')
    if payment_proof:
        payment_text = payment_proof.get('text') if isinstance(payment_proof, dict) else payment_proof
    else:
        payment_text = None

    if payment_text:
        _add_body_text(doc, payment_text)
    else:
        _add_body_text(doc,
            "The following payment details confirm that the transaction was successfully authorised and "
            "captured. No credit or refund was issued against this payment."
        )
    _add_image(doc, screenshots.get('card_details_screenshot'), "Payment Gateway Card Details")

    # Section 3: Shipping Proof
    _add_section_header(doc, "3. Shipping proof")
    shipping_text = "Please find below the Delivery Confirmation Receipt confirming successful delivery of the order:"
    if data.get('shipping_proof'):
        if isinstance(data['shipping_proof'], dict):
            shipping_text = data['shipping_proof'].get('text', shipping_text)
    _add_body_text(doc, shipping_text)
    _add_image(doc, screenshots.get('tracking_screenshot'), "Carrier Delivery Confirmation")

    # Section 4: Return Policy
    _add_section_header(doc, "4. Merchant's Return/Exchange Policy")
    _add_body_text(doc, return_policy["text"])
    if return_policy.get("url"):
        _add_body_text(doc, "Merchant's Returns & Exchanges Policy can be found by the below link:")
        _add_link_text(doc, return_policy['url'])
    _add_body_text(doc, "Here is an extract from merchant's refund policy:")
    _add_image(doc, return_policy_image, "Return Policy Screenshot")

    # Closing
    _add_thin_rule(doc)
    _add_section_header(doc, "Summary")
    if data.get('closing_statement'):
        _add_conclusion_text(doc, data['closing_statement'])
    else:
        _add_conclusion_text(doc,
            "The transaction was successfully authorised, captured, and the order was delivered to the "
            "cardholder as confirmed by the shipping documentation above. No credit or refund was agreed "
            "upon or processed. The cardholder did not attempt to return the product in accordance with the "
            "merchant's Returns & Exchanges Policy. These facts confirm that this chargeback is unwarranted "
            "and should be cancelled."
        )

    _add_footer(doc)
    doc.save(output_path)
    print(f"DOCX created: {output_path}")
    return output_path
